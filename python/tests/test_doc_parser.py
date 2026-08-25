"""
文档解析层测试 — 覆盖 PDF / DOCX / Markdown / 表格 / 分块元数据

运行: pytest tests/test_doc_parser.py -v
注意: marker-pdf / docling 未安装或首次下载模型时, 自动走 pymupdf 降级路径
"""

from __future__ import annotations

import os
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from agents.doc_parser_agent import DocParserAgent, DocType  # noqa: E402


def _make_tmp(tmp_path: Path, name: str, content: bytes) -> str:
    p = tmp_path / name
    p.write_bytes(content)
    return str(p)


# ── 固定分块路径 (不依赖 embedding 模型) ──────────────────────

@pytest.fixture
def parser_structural(monkeypatch) -> DocParserAgent:
    from config import settings
    monkeypatch.setattr(settings, "semantic_chunk_enabled", False)
    monkeypatch.setattr(settings, "marker_enabled", False)
    return DocParserAgent(embeddings=None)


@pytest.mark.asyncio
async def test_parse_markdown_with_heading_path(parser_structural, tmp_path):
    """Markdown 解析: 标题层级 → heading_path 元数据"""
    md = """# 公司简介

AgentKnowledgeHub 是一个企业级知识管理系统。

## 团队架构

### 技术部
张三负责 AI 平台研发。

## 产品介绍
支持多模态文档解析。
"""
    path = _make_tmp(tmp_path, "test.md", md.encode("utf-8"))
    chunks = await parser_structural.parse(path)

    assert chunks, "Markdown 解析不应为空"
    # 每个 chunk 都带 heading_path 元数据
    for c in chunks:
        assert "heading_path" in c.metadata

    # 技术部段落应携带完整标题路径
    tech_chunk = next(c for c in chunks if "张三" in c.content)
    heading = tech_chunk.metadata["heading_path"]
    assert "团队架构" in heading, f"heading_path 缺失父级标题: {heading}"
    assert "技术部" in heading, f"heading_path 缺失当前标题: {heading}"
    assert "产品介绍" not in heading, "heading_path 不应包含无关标题"


@pytest.mark.asyncio
async def test_parse_markdown_table(parser_structural, tmp_path):
    """Markdown 表格 → is_table 元数据 + 结构化内容"""
    md = """# 财务报表

| 季度 | 营收 | 利润 |
|------|------|------|
| Q1   | 100  | 20   |
| Q2   | 150  | 30   |
"""
    path = _make_tmp(tmp_path, "table.md", md.encode("utf-8"))
    chunks = await parser_structural.parse(path)

    table_chunks = [c for c in chunks if c.metadata.get("is_table")]
    assert table_chunks, "应识别出表格 chunk"
    assert "Q1" in table_chunks[0].content
    assert "季度" in table_chunks[0].content, "表格表头应保留"


@pytest.mark.asyncio
async def test_parse_pdf_with_pymupdf(parser_structural, tmp_path):
    """PDF 解析: pymupdf 提取文本与表格"""
    import fitz  # pymupdf

    pdf_path = tmp_path / "test.pdf"
    doc = fitz.open()
    page = doc.new_page()
    # 默认 Helvetica 不支持中文, 使用 pymupdf 内置中文字体 china-s
    page.insert_text((72, 72), "AgentKnowledgeHub 年度报告", fontsize=16, fontname="china-s")  # 大字号 → 标题
    page.insert_text((72, 120), "本年度营收增长 30%，新增客户 120 家。", fontsize=11, fontname="china-s")
    page.insert_text((72, 160), "张三担任技术总监。", fontsize=11, fontname="china-s")
    doc.save(str(pdf_path))
    doc.close()

    chunks = await parser_structural.parse(str(pdf_path))
    assert chunks, "PDF 解析不应为空"

    full_text = "\n".join(c.content for c in chunks)
    assert "AgentKnowledgeHub" in full_text, "PDF 文本提取失败"
    assert "张三" in full_text

    # 字号启发式: 16px 应识别为标题
    headings = [c for c in chunks if c.metadata.get("heading_path")]
    assert headings, "大字号文本应识别为标题"


@pytest.mark.asyncio
async def test_parse_docx(parser_structural, tmp_path):
    """DOCX 解析: 段落 + 标题 + 表格"""
    import docx as docx_lib

    path = tmp_path / "test.docx"
    d = docx_lib.Document()
    d.add_heading("员工手册", level=1)
    d.add_heading("考勤制度", level=2)
    d.add_paragraph("员工需在 9:00 前到岗。")
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "姓名"
    table.rows[0].cells[1].text = "部门"
    table.rows[1].cells[0].text = "李四"
    table.rows[1].cells[1].text = "市场部"
    d.save(str(path))

    chunks = await parser_structural.parse(str(path))
    full_text = "\n".join(c.content for c in chunks)
    assert "员工手册" in full_text, "DOCX 标题解析失败"
    assert "考勤制度" in full_text
    assert "李四" in full_text, "DOCX 表格解析失败"

    table_chunks = [c for c in chunks if c.metadata.get("is_table")]
    assert table_chunks, "DOCX 表格应标记 is_table"


@pytest.mark.asyncio
async def test_parse_excel(parser_structural, tmp_path):
    """Excel 解析: 多 sheet 结构化"""
    import openpyxl

    path = tmp_path / "test.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "员工列表"
    ws.append(["姓名", "职位"])
    ws.append(["张三", "CEO"])
    wb.save(str(path))

    chunks = await parser_structural.parse(str(path))
    assert chunks
    assert any("张三" in c.content for c in chunks), "Excel 内容解析失败"
    assert all(c.metadata.get("is_table") for c in chunks), "Excel chunk 应标记 is_table"


@pytest.mark.asyncio
async def test_fixed_chunk_fallback(parser_structural, tmp_path):
    """长文本: 超长段落按结构边界切分, 不切断句子"""
    sentence = "这是用于测试分块边界的一段较长的句子。" * 60
    path = _make_tmp(tmp_path, "long.txt", sentence.encode("utf-8"))
    chunks = await parser_structural.parse(path)

    assert len(chunks) > 1, "长文本应被切分为多个 chunk"
    for c in chunks:
        assert len(c.content) <= 512 + 64, f"chunk 超长: {len(c.content)}"
    # 结构分块应优先按句子边界切, chunk 结尾应为完整句子
    for c in chunks[:-1]:
        assert c.content.rstrip().endswith("。") or c.content.rstrip().endswith("；"), \
            f"chunk 在句子中间被切断: ...{c.content[-20:]}"


@pytest.mark.asyncio
async def test_unsupported_file(tmp_path):
    """不支持的文件类型: 不崩溃, 返回空"""
    parser = DocParserAgent(embeddings=None)
    path = _make_tmp(tmp_path, "data.bin", b"\x00\x01\x02")
    chunks = await parser.parse(path)
    assert isinstance(chunks, list)


@pytest.mark.asyncio
async def test_classify_extensions():
    """扩展名分类: 覆盖所有支持格式"""
    cases = {
        "a.pdf": DocType.PDF,
        "a.docx": DocType.WORD,
        "a.doc": DocType.WORD,
        "a.pptx": DocType.PPT,
        "a.png": DocType.IMAGE,
        "a.csv": DocType.TABLE,
        "a.xlsx": DocType.TABLE,
        "a.txt": DocType.TEXT,
        "a.md": DocType.MARKDOWN,
    }
    agent = DocParserAgent(embeddings=None)
    for fname, expected in cases.items():
        assert agent._classify(fname) == expected, f"{fname} 分类错误"


# ── 语义分块 (用确定性 fake embeddings, 不下载模型) ────────────

class FakeEmbeddings:
    """确定性 fake embeddings: 相同文本 → 相同向量"""

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [self._vec(t) for t in texts]

    def embed_query(self, text: str) -> list[float]:
        return self._vec(text)

    async def aembed_documents(self, texts: list[str]) -> list[list[float]]:
        return self.embed_documents(texts)

    async def aembed_query(self, text: str) -> list[float]:
        return self.embed_query(text)

    @staticmethod
    def _vec(text: str) -> list[float]:
        return [float(len(text) % 7), float(hash(text) % 11)]


def test_semantic_chunking_with_fake_embeddings():
    """语义分块器可注入 embeddings, 不依赖真实模型"""
    from services.document_parser import ParseBlock
    from utils.chunker import DocumentChunker

    blocks = [
        ParseBlock(kind="paragraph", content="段落A: 机器学习的基础知识介绍。" * 5),
        ParseBlock(kind="paragraph", content="段落B: 完全不同的主题内容。" * 5),
    ]
    chunker = DocumentChunker(embeddings=FakeEmbeddings(), semantic_enabled=True)
    chunks = chunker.chunk(blocks, doc_id="d1", doc_type=DocType.TEXT, source="test")

    assert chunks, "语义分块不应为空"
    assert all("heading_path" in c.metadata for c in chunks)


def test_semantic_chunker_degrades_gracefully():
    """无 embeddings 时自动降级为结构分块, 不抛异常"""
    from services.document_parser import ParseBlock
    from utils.chunker import DocumentChunker

    blocks = [ParseBlock(kind="paragraph", content="一段普通文本。" * 100)]
    chunker = DocumentChunker(embeddings=None, semantic_enabled=True)
    chunks = chunker.chunk(blocks, doc_id="d1", doc_type=DocType.TEXT, source="test")
    assert chunks
